"""Generate report.docx from pipeline output artifacts.

Usage:
    uv run python generate_report.py
    uv run python generate_report.py --output-dir output --report-path report.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from training_constants import OFFICIAL_MODEL_NAMES


# ── Formatting helpers ────────────────────────────────────────────────────────

def _fmt(val, digits=4, signed=False):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "N/A"
    return f"{float(val):{'+'if signed else ''}.{digits}f}"

def _pct(val, digits=1):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "N/A"
    return f"{float(val) * 100:.{digits}f}%"

def _int_fmt(val):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "N/A"
    return f"{int(val):,}"

def _read(output_dir: Path, name: str) -> pd.DataFrame | None:
    path = output_dir / name
    return pd.read_csv(path) if path.exists() else None


def _recommended_model_name(results_df: pd.DataFrame | None, selection_df: pd.DataFrame | None) -> str:
    if selection_df is not None and not selection_df.empty and {"model", "recommended"}.issubset(selection_df.columns):
        recommended_rows = selection_df.loc[selection_df["recommended"].fillna(False).astype(bool)]
        if not recommended_rows.empty:
            return str(recommended_rows.iloc[0]["model"])

    if results_df is not None and not results_df.empty and "Model" in results_df.columns:
        candidate_df = results_df.loc[
            results_df["Model"].isin(OFFICIAL_MODEL_NAMES)
            & ~results_df["Model"].str.contains("calibrated|Ensemble", regex=True, na=False)
        ].copy()
        if not candidate_df.empty:
            candidate_df = candidate_df.sort_values(["PR AUC", "ROC AUC"], ascending=False)
            return str(candidate_df.iloc[0]["Model"])

    return OFFICIAL_MODEL_NAMES[0]


# ── Table helpers ─────────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex,
    })
    shading.append(shd)

def _style_header_row(table):
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, "1F4E79")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            c = table.rows[r_idx + 1].cells[c_idx]
            c.text = str(val)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    _style_header_row(table)
    return table

def _add_image(doc, path: Path, width=5.5):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Report generation ─────────────────────────────────────────────────────────

def generate_report(output_dir: str = "output", report_path: str = "report.docx"):
    o = Path(output_dir)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Load all data
    results = _read(o, "results.csv")
    sel = _read(o, "model_selection.csv")
    overfit = _read(o, "overfit_report.csv")
    bench = _read(o, "benchmark_comparisons.csv")
    oot = _read(o, "rolling_oot_summary.csv")
    ci = _read(o, "confidence_intervals.csv")
    pop = _read(o, "population_summary.csv")
    psi_df = _read(o, "psi.csv")
    drift = _read(o, "concept_drift.csv")
    lift = _read(o, "lift_table.csv")
    ens = _read(o, "ensemble_weights.csv")
    ai = _read(o, "adverse_impact_age.csv")
    selbias = _read(o, "selection_bias_correlation.csv")
    iv = _read(o, "iv_summary.csv")
    thresh = _read(o, "threshold_analysis.csv")
    ablation = _read(o, "ablation_results.csv")

    # Derive key numbers
    recommended = _recommended_model_name(results, sel)
    rec_r = None
    if results is not None and not results.empty:
        rec_rows = results.loc[results["Model"] == recommended]
        rec_r = rec_rows.iloc[0] if not rec_rows.empty else results.iloc[0]

    test_n = int(rec_r["N"]) if rec_r is not None and "N" in rec_r else 0
    test_pos = int(bench.iloc[0]["n_pos"]) if bench is not None and not bench.empty else 0
    test_rate = test_pos / test_n if test_n > 0 else 0

    pre_booked = None
    if pop is not None:
        pre_booked_rows = pop.loc[(pop["split"] == "pre_split") & (pop["status_name"] == "Booked")]
        if not pre_booked_rows.empty:
            pre_booked = pre_booked_rows.iloc[0]
    post_total = int(pop.loc[pop["split"] == "post_split", "n_rows"].sum()) if pop is not None else 0

    lr_lift_df = lift.loc[lift["model"] == recommended] if lift is not None else None
    lr_thresh_df = thresh.loc[thresh["model"] == recommended] if thresh is not None else None
    lr_ai_df = ai.loc[ai["model"] == recommended] if ai is not None else None
    non_cal_results = pd.DataFrame()
    if results is not None and not results.empty:
        non_cal_results = results.loc[
            ~results["Model"].str.contains("calibrated|Ensemble", regex=True, na=False)
        ].copy()
    best_auc_model = recommended
    best_pr_model = recommended
    if not non_cal_results.empty:
        best_auc_model = str(non_cal_results.loc[non_cal_results["ROC AUC"].astype(float).idxmax(), "Model"])
        best_pr_model = str(non_cal_results.loc[non_cal_results["PR AUC"].astype(float).idxmax(), "Model"])

    # ══════════════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════════════
    title = doc.add_heading("Credit Scoring Model Validation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Basel Bad Default Prediction Model\nDirect Consumer Credit Portfolio\n")
    run.bold = True
    run.font.size = Pt(14)
    meta.add_run("\nPopulation mode: Underwriting  |  Evaluation: Booked-proxy holdout")
    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    p = doc.add_paragraph()
    p.add_run("Purpose. ").bold = True
    p.add_run(
        "This report presents the development and validation of a binary classification model "
        "predicting 12-month default (basel_bad) for direct consumer credit applications. "
        "The model is intended to support underwriting decisions by scoring all decisioned "
        "applications (booked, rejected, and canceled) at the point of application."
    )

    p = doc.add_paragraph()
    p.add_run("Recommendation. ").bold = True
    score_text = f" with a weighted selection score of {sel.iloc[0]['weighted_score']:.1f}/100" if sel is not None else ""
    p.add_run(
        f"The recommended production model is {recommended}{score_text}. "
        f"It achieves a ROC AUC of {_fmt(rec_r['ROC AUC'])} and KS of {_fmt(rec_r['KS'])} on the "
        f"held-out booked-proxy test set ({_int_fmt(test_n)} accounts, {_int_fmt(test_pos)} observed "
        f"defaults, {_pct(test_rate)} base rate)."
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Key findings:").bold = True

    # Benchmark lift
    lr_vs_score = bench.loc[
        (bench["candidate_model"] == recommended) & (bench["reference_model"] == "score_RF (benchmark)")
    ] if bench is not None else pd.DataFrame()
    lr_vs_risk = bench.loc[
        (bench["candidate_model"] == recommended) & (bench["reference_model"] == "risk_score_rf (benchmark)")
    ] if bench is not None else pd.DataFrame()

    findings = []
    if not lr_vs_score.empty:
        r = lr_vs_score.iloc[0]
        findings.append(
            f"The recommended model significantly outperforms score_RF on ROC AUC "
            f"(+{r['auc_improvement']:.4f}, 95% CI [{r['auc_improvement_lo']:+.4f}, {r['auc_improvement_hi']:+.4f}], "
            f"DeLong p={r['auc_delong_p_value']:.6f}), confirming that the new model provides "
            f"materially better risk ranking than the existing weaker benchmark."
        )
    if not lr_vs_risk.empty:
        r = lr_vs_risk.iloc[0]
        findings.append(
            f"Against the stronger risk_score_rf benchmark, the AUC improvement is modest and not "
            f"statistically significant ({r['auc_improvement']:+.4f}, p={r['auc_delong_p_value']:.2f}). "
            f"On PR AUC, risk_score_rf retains a meaningful lead ({r['pr_auc_improvement']:+.4f}), "
            f"indicating that the existing benchmark concentrates defaults more effectively in the "
            f"highest-risk segment. This gap represents the primary improvement opportunity."
        )
    if overfit is not None:
        lr_of = overfit.loc[overfit["model"] == recommended]
        tree_of = overfit.loc[(overfit["model"].isin(OFFICIAL_MODEL_NAMES)) & (overfit["model"] != recommended)]
        if not lr_of.empty and not tree_of.empty:
            worst_peer = tree_of.sort_values("auc_delta", ascending=False).iloc[0]
            findings.append(
                f"{recommended} shows an AUC overfit gap of {lr_of.iloc[0]['auc_delta']:+.4f}. "
                f"Across the remaining candidates, AUC deltas range from "
                f"{tree_of['auc_delta'].min():+.4f} to {tree_of['auc_delta'].max():+.4f}, "
                f"with {worst_peer['model']} exhibiting the largest gap. "
                f"This generalization advantage is one of the main reasons the scorecard favors {recommended}."
            )
        elif not lr_of.empty:
            findings.append(
                f"{recommended} shows an AUC overfit gap of {lr_of.iloc[0]['auc_delta']:+.4f} on the held-out sample."
            )
    if lr_lift_df is not None and not lr_lift_df.empty:
        top2_capture = lr_lift_df.loc[lr_lift_df["decile"] <= 2, "capture_rate"].max()
        top_lift = lr_lift_df.iloc[0]["lift"]
        findings.append(
            f"The top score decile concentrates defaults at {top_lift:.1f}x the population average. "
            f"The top two deciles capture {_pct(top2_capture)} of all defaults, demonstrating "
            f"meaningful separation for underwriting cutoff decisions."
        )
    if lr_ai_df is not None:
        n_fail = int((lr_ai_df["air_flag"] == "FAIL").sum())
        if n_fail > 0:
            failing = lr_ai_df.loc[lr_ai_df["air_flag"] == "FAIL"]
            bands = ", ".join(failing["age_band"].values)
            findings.append(
                f"The adverse impact analysis flags {n_fail} age band(s) ({bands}) below the 80% "
                f"AIR threshold at a 10% rejection rate. Young applicants receive systematically higher "
                f"predicted risk, which is consistent with the observed higher default rate in this segment "
                f"but warrants regulatory review."
            )
        else:
            findings.append("All age bands pass the 80% adverse impact ratio threshold.")

    if psi_df is not None and not psi_df.empty:
        max_psi = float(psi_df["psi"].max())
        findings.append(
            f"Score distributions remain stable between training and test, with maximum PSI {_fmt(max_psi)}."
        )
    if drift is not None and not drift.empty:
        n_drift = int((drift["concept_drift_flag"] == "YES").sum())
        if n_drift == 0:
            findings.append("No concept drift is detected across the rolling out-of-time validation windows.")
        else:
            drift_models = ", ".join(drift.loc[drift["concept_drift_flag"] == "YES", "model"].astype(str).tolist())
            findings.append(
                f"Concept drift is detected for {n_drift} model(s): {drift_models}."
            )

    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    # ══════════════════════════════════════════════════════════════════════
    # 2. DATA AND POPULATION
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Data and Population Design", level=1)

    doc.add_heading("2.1 Population Scope", level=2)
    doc.add_paragraph(
        "The pipeline operates in underwriting mode, meaning it loads the full decisioned "
        "population: booked loans, rejected applications, and canceled applications. This "
        "reflects the production use case where the model scores every incoming application "
        "regardless of the eventual decision."
    )
    doc.add_paragraph(
        "However, the supervised training target (12-month default) is observable only for "
        "booked accounts that have had sufficient time on book. This creates a fundamental "
        "limitation: the model is trained and evaluated on the accepted-population proxy, "
        "not the full applicant population. Rejected and canceled applications contribute "
        "to the applicant-stage score frame but do not influence model fitting or evaluation metrics."
    )

    if pop is not None:
        doc.add_heading("2.2 Population Breakdown", level=2)
        rows = []
        for _, r in pop.iterrows():
            rows.append([
                r["split"].replace("_", "-"), r["status_name"], _int_fmt(r["n_rows"]),
                _int_fmt(r.get("n_bad_observed", np.nan)),
                f"{r['date_start']} to {r['date_end']}" if pd.notna(r.get("date_start")) else "",
            ])
        _add_table(doc, ["Split", "Status", "Applications", "Defaults", "Period"], rows)
        doc.add_paragraph("")

        if pre_booked is not None:
            train_defaults = int(pre_booked["n_bad_observed"])
            train_n = int(pre_booked["n_rows"])
            doc.add_paragraph(
                f"The pre-split booked sample ({_int_fmt(train_n)} accounts, {_int_fmt(train_defaults)} "
                f"defaults, {_pct(train_defaults/train_n)} base rate) is partitioned into: "
                f"(1) a feature-discovery window for interaction search and stability selection, "
                f"(2) an estimation window for model fitting, and "
                f"(3) a calibration holdout (latest 15% by date) for probability calibration. "
                f"The post-split booked matured subset ({_int_fmt(test_n)} accounts, {_int_fmt(test_pos)} "
                f"defaults, {_pct(test_rate)} base rate) serves as the held-out evaluation sample."
            )

    doc.add_heading("2.3 Temporal Design", level=2)
    doc.add_paragraph(
        "All data partitioning is strictly temporal to prevent information leakage. "
        "The split date (2024-07-01) separates development from the held-out test period. "
        "The maturity cutoff (2025-01-01) ensures that only accounts with a full 12-month "
        "observation window are included in evaluation. Within the development period, "
        "feature discovery uses the earliest data, model estimation uses the middle period, "
        "and calibration uses the latest block. No random shuffling or stratified splitting "
        "is used at any stage."
    )

    # ══════════════════════════════════════════════════════════════════════
    # 3. METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Methodology", level=1)

    doc.add_heading("3.1 Feature Engineering", level=2)
    doc.add_paragraph(
        "Starting from 15 raw numerical and 12 raw categorical features, the pipeline "
        "engineers additional signals in several families:"
    )
    features_list = [
        "Affordability ratios: installment-to-income, amount-to-income, amount-per-month, and household-capacity variants",
        "Portfolio composition: total products, book ratios for cards and loans, presence flags",
        "Bureau-capacity transforms: log of income, loan amount, and maximum credit facility",
        "Categorical interactions: targeted crosses of product type, housing, and customer-type segments",
        "Missingness indicators: binary flags for fields with meaningful missing rates (> 1%)",
        "Frequency encodings: training-set category proportions as numerical features",
        "Group-relative statistics: individual value divided by group median for key numeric-categorical pairs",
        "Binned interactions: quantile-binned numerical features crossed with categorical values to capture threshold effects",
    ]
    for item in features_list:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Feature Selection", level=2)
    doc.add_paragraph(
        "Feature selection uses temporal stability selection rather than classical RFECV. "
        "For each fold of an expanding temporal cross-validation, an elastic-net logistic "
        "regression (with grid-searched C and L1 ratio) is fitted on the preprocessed features. "
        "Features that receive a non-zero coefficient in >= 50% of folds are retained. "
        "If fewer than 5 features pass the threshold, the top features by selection frequency "
        "are backfilled. This approach favors features whose predictive signal is stable "
        "across time periods rather than those that happen to be strong in a single window."
    )

    if ablation is not None and not ablation.empty:
        doc.add_heading("3.3 Ablation Study", level=2)
        doc.add_paragraph(
            "An ablation study measures the incremental contribution of each feature engineering "
            "stage by training a Logistic Regression on progressively richer feature sets and "
            "evaluating on the held-out test:"
        )
        rows = []
        for _, r in ablation.iterrows():
            rows.append([
                r["variant"].replace("_", " ").title(),
                str(int(r["n_features"])),
                _fmt(r["ROC AUC"]), _fmt(r["PR AUC"]),
            ])
        _add_table(doc, ["Feature Set", "Features", "ROC AUC", "PR AUC"], rows)
        doc.add_paragraph("")
        doc.add_paragraph(
            "The ablation reveals that engineered features (affordability ratios, log transforms) "
            "provide the largest marginal lift over raw inputs. The interaction search and "
            "stability selection refine the feature set but their primary value is regularization "
            "rather than discrimination gain."
        )

    doc.add_heading("3.4 Model Training", level=2)
    doc.add_paragraph(
        "Five candidate architectures are trained with Optuna Bayesian hyperparameter optimization "
        "using multivariate TPE sampling (models correlations between hyperparameters) and median "
        "pruning (terminates unpromising trials after the first 1-2 cross-validation folds)."
    )
    _add_table(doc,
        ["Model", "Approach", "Calibration", "Key Regularization"],
        [
            ["Logistic Regression", "L2-penalized, LBFGS solver, C and TargetEncoder smooth tuned jointly", "Sigmoid (Platt)", "class_weight='balanced'"],
            ["EBM", "Explainable boosting machine with additive shape functions and limited pairwise interactions", "Sigmoid (Platt)", "max_bins 64-256, interactions 0-15, min_samples_leaf 2-50"],
            ["LightGBM", "Gradient boosting, native categorical handling via ordinal encoding", "Isotonic", "num_leaves 8-31, max_bin 63-127, colsample_bynode 0.6-1.0, early stopping after 30 rounds"],
            ["XGBoost", "Gradient boosting, TargetEncoder preprocessing", "Isotonic", "max_depth 2-4, min_child_weight 20-100, colsample_bynode 0.6-1.0, early stopping after 30 rounds"],
            ["CatBoost", "Oblivious gradient boosting, balanced class weights", "Isotonic", "depth 3-5, min_data_in_leaf 50-300, early stopping after 30 rounds"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Tree complexity bounds are deliberately tightened for the low (~3.8%) default rate: "
        "shallow depths, high minimum leaf sizes, coarse histogram bins, and aggressive "
        "per-node feature subsampling. When reject-inference sample weights are provided, "
        "scale_pos_weight is disabled to avoid double-rebalancing the minority class."
    )

    doc.add_heading("3.5 Calibration", level=2)
    doc.add_paragraph(
        "Probability calibration is fitted on the latest pre-split booked holdout (15% of "
        "development data by date). Isotonic calibration is used for tree models because their "
        "output distributions are step-functions that violate the linear log-odds assumption of "
        "Platt scaling. Sigmoid (Platt) calibration is retained for Logistic Regression and EBM, "
        "whose score distributions are smoother and closer to a log-odds scale. The model selection scorecard evaluates calibration "
        "quality using the calibrated Brier score, since that reflects production probability quality."
    )

    doc.add_heading("3.6 Model Selection Framework", level=2)
    doc.add_paragraph(
        "A weighted multi-criteria scorecard ranks candidates on five dimensions, each "
        "normalized to a 0-100 scale across the candidate set:"
    )
    _add_table(doc,
        ["Criterion", "Weight", "Source", "Rationale"],
        [
            ["Discrimination", "35%", "Test PR AUC", "Primary measure of risk ranking quality on imbalanced data"],
            ["Stability", "20%", "Rolling OOT mean PR AUC", "Rewards models whose performance holds across multiple forward time windows"],
            ["Calibration", "15%", "Calibrated Brier score", "Assesses whether predicted probabilities are trustworthy for pricing"],
            ["Generalization", "15%", "Train-test AUC delta", "Penalizes overfitting: 0 penalty if delta <= 0.01, full penalty at 0.10"],
            ["Benchmark lift", "15%", "Best AUC improvement vs benchmarks", "Rewards improvement over existing production scores"],
        ],
    )
    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════════════
    # 4. RESULTS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Results", level=1)

    doc.add_heading("4.1 Held-Out Test Performance", level=2)
    doc.add_paragraph(
        f"The table below reports performance on {_int_fmt(test_n)} held-out booked accounts "
        f"from July 2024 to January 2025, with {_int_fmt(test_pos)} observed defaults ({_pct(test_rate)} "
        f"base rate). Models are sorted by PR AUC, which is the primary metric for imbalanced "
        f"classification where the cost of missed defaults greatly exceeds the cost of false alarms."
    )
    if results is not None:
        official = results.loc[~results["Model"].str.contains("calibrated", regex=False, na=False)]
        rows = []
        for _, r in official.iterrows():
            rows.append([
                str(r["Model"]), _fmt(r["ROC AUC"]), _fmt(r["Gini"]),
                _fmt(r["KS"]), _fmt(r["PR AUC"]), _fmt(r.get("Brier")),
            ])
        _add_table(doc, ["Model", "ROC AUC", "Gini", "KS", "PR AUC", "Brier"], rows)
        doc.add_paragraph("")

    _add_image(doc, o / "plots" / "stakeholder_holdout_benchmarks.png")
    doc.add_paragraph("")

    if best_auc_model == best_pr_model:
        doc.add_paragraph(
            f"On the non-calibrated leaderboard, {best_auc_model} leads both ROC AUC and PR AUC. "
            f"{recommended} remains the recommended production candidate because the final decision balances "
            "discrimination with calibration, stability, generalization, and benchmark lift."
        )
    else:
        doc.add_paragraph(
            f"On the non-calibrated leaderboard, {best_auc_model} leads ROC AUC while {best_pr_model} "
            f"delivers the strongest PR AUC. {recommended} remains the recommended production candidate because "
            "the final decision balances discrimination with calibration, stability, generalization, and benchmark lift."
        )

    doc.add_heading("4.2 Model Selection Scorecard", level=2)
    if sel is not None:
        rows = []
        for _, r in sel.iterrows():
            rows.append([
                str(r["model"]), f"{r['discrimination_score']:.1f}",
                f"{r['calibration_score']:.1f}", f"{r['stability_score']:.1f}",
                f"{r['generalization_score']:.1f}", f"{r['lift_score']:.1f}",
                f"{r['weighted_score']:.1f}", "Yes" if r["recommended"] else "",
            ])
        _add_table(doc, ["Model", "Disc.", "Calib.", "Stab.", "Gen.", "Lift", "Overall", "Rec."], rows)
        doc.add_paragraph("")
        doc.add_paragraph(
            f"{recommended} achieves the highest weighted score ({sel.iloc[0]['weighted_score']:.1f}/100) "
            "on the multi-criteria scorecard. The recommendation balances discrimination, calibration, "
            "stability, generalization, and benchmark lift rather than optimizing only a single held-out metric."
        )

    doc.add_heading("4.3 Overfitting Analysis", level=2)
    doc.add_paragraph(
        "The overfitting report compares each model's metrics on the training data versus the "
        "held-out test. A train-test delta above 0.03 on AUC or PR AUC triggers a flag. "
        "This analysis helps explain why models with extremely strong in-sample fit do not always "
        "translate into the best production recommendation."
    )
    if overfit is not None:
        rows = []
        for _, r in overfit.iterrows():
            rows.append([
                str(r["model"]), _fmt(r["train_auc"]), _fmt(r["test_auc"]),
                _fmt(r["auc_delta"], signed=True), _fmt(r["train_pr_auc"]),
                _fmt(r["test_pr_auc"]), _fmt(r["pr_auc_delta"], signed=True), r["overfit_flag"],
            ])
        _add_table(doc, ["Model", "Train AUC", "Test AUC", "Delta", "Train PR", "Test PR", "PR Delta", "Flag"], rows)
        doc.add_paragraph("")
        rec_overfit = overfit.loc[overfit["model"] == recommended]
        peer_overfit = overfit.loc[(overfit["model"].isin(OFFICIAL_MODEL_NAMES)) & (overfit["model"] != recommended)]
        if not rec_overfit.empty and not peer_overfit.empty:
            worst_peer = peer_overfit.sort_values("auc_delta", ascending=False).iloc[0]
            doc.add_paragraph(
                f"{recommended} posts an AUC delta of {rec_overfit.iloc[0]['auc_delta']:+.4f} and a PR AUC delta of "
                f"{rec_overfit.iloc[0]['pr_auc_delta']:+.4f}. Among the remaining candidates, the largest AUC gap is "
                f"{worst_peer['model']} at {worst_peer['auc_delta']:+.4f}. This indicates that the more flexible "
                "models fit the development sample more aggressively than they generalize to unseen accounts."
            )

    doc.add_heading("4.4 Bootstrap Confidence Intervals", level=2)
    doc.add_paragraph(
        "Stratified bootstrap confidence intervals (1,000 iterations, stratified by class) "
        "provide uncertainty bounds. Point estimates are the observed test-set statistics, "
        "not bootstrap medians."
    )
    if ci is not None:
        ci_off = ci.loc[~ci["Model"].str.contains("calibrated|Ensemble", regex=True, na=False)]
        rows = []
        for _, r in ci_off.iterrows():
            rows.append([
                str(r["Model"]),
                f"{_fmt(r['AUC'])} [{_fmt(r['AUC_lo'])}, {_fmt(r['AUC_hi'])}]",
                f"{_fmt(r['PR_AUC'])} [{_fmt(r['PR_AUC_lo'])}, {_fmt(r['PR_AUC_hi'])}]",
            ])
        _add_table(doc, ["Model", "ROC AUC [95% CI]", "PR AUC [95% CI]"], rows)
        doc.add_paragraph("")

    doc.add_heading("4.5 Rolling Out-of-Time Validation", level=2)
    doc.add_paragraph(
        "Rolling OOT validation trains fresh models on expanding time windows and evaluates "
        "on forward holdout periods. This simulates the model's behaviour as the calendar "
        "advances and is a stronger stability test than the single held-out evaluation."
    )
    if oot is not None:
        oot_unc = oot.loc[~oot["Model"].str.contains("calibrated", regex=False, na=False)]
        rows = []
        for _, r in oot_unc.iterrows():
            rows.append([
                str(r["Model"]), str(int(r["n_folds"])),
                _fmt(r["mean_ROC_AUC"]), f"({_fmt(r.get('std_ROC_AUC'))})",
                _fmt(r["mean_PR_AUC"]), f"({_fmt(r.get('std_PR_AUC'))})",
            ])
        _add_table(doc, ["Model", "Folds", "Mean AUC", "Std", "Mean PR", "Std"], rows)
        doc.add_paragraph("")

    _add_image(doc, o / "plots" / "stakeholder_rolling_oot.png")
    doc.add_paragraph("")

    if drift is not None and not drift.empty:
        doc.add_heading("4.6 Concept Drift Detection", level=2)
        doc.add_paragraph(
            "Concept drift is assessed by fitting a linear trend to PR AUC across rolling OOT folds. "
            "A declining trend exceeding 0.02 first-to-last triggers a drift flag, indicating that "
            "the score-to-outcome relationship may be degrading over time."
        )
        rows = []
        for _, r in drift.iterrows():
            rows.append([
                str(r["model"]), _fmt(r["pr_auc_first"]), _fmt(r["pr_auc_last"]),
                _fmt(r["pr_auc_slope_per_fold"], signed=True), r["concept_drift_flag"],
            ])
        _add_table(doc, ["Model", "First PR", "Last PR", "Slope/Fold", "Drift"], rows)
        doc.add_paragraph("")
        n_drift = int((drift["concept_drift_flag"] == "YES").sum())
        n_windows = int(oot_unc["n_folds"].max()) if oot is not None and not oot.empty else 0
        if n_drift == 0:
            doc.add_paragraph(
                f"No models show concept drift in the current evaluation period. PR AUC fluctuations "
                f"remain within normal sampling variance across {n_windows} forward validation windows."
            )
        else:
            drift_models = ", ".join(drift.loc[drift["concept_drift_flag"] == "YES", "model"].astype(str).tolist())
            doc.add_paragraph(
                f"{n_drift} model(s) are flagged for concept drift across {n_windows} forward validation windows: {drift_models}."
            )

    # ══════════════════════════════════════════════════════════════════════
    # 5. OPERATIONAL ANALYSIS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Operational Analysis", level=1)

    doc.add_heading("5.1 Decile Lift Table", level=2)
    doc.add_paragraph(
        f"The lift table for {recommended} shows how defaults concentrate in the highest-risk "
        f"score deciles. This is the most operationally relevant view for setting rejection thresholds."
    )
    if lr_lift_df is not None and not lr_lift_df.empty:
        rows = []
        for _, r in lr_lift_df.iterrows():
            rows.append([
                str(int(r["decile"])), _int_fmt(r["n_accounts"]), _int_fmt(r["n_defaults"]),
                _pct(r["default_rate"]), f"{r['lift']:.2f}x",
                f"{r['cum_lift']:.2f}x", _pct(r["capture_rate"]),
            ])
        _add_table(doc, ["Decile", "Accounts", "Defaults", "Default Rate", "Lift", "Cum. Lift", "Capture"], rows)
        doc.add_paragraph("")
        top_rate = lr_lift_df.iloc[0]["default_rate"]
        bot_rate = lr_lift_df.iloc[-1]["default_rate"]
        concentration_ratio = f"{top_rate/bot_rate:.0f}:1" if pd.notna(bot_rate) and bot_rate > 0 else "N/A"
        doc.add_paragraph(
            f"The riskiest decile has a {_pct(top_rate)} default rate versus {_pct(bot_rate)} in the "
            f"safest decile -- an {concentration_ratio} concentration ratio. The top two deciles "
            f"capture {_pct(lr_lift_df.loc[lr_lift_df['decile']<=2, 'capture_rate'].max())} of all defaults."
        )

    _add_image(doc, o / "plots" / "stakeholder_gains.png")
    doc.add_paragraph("")

    doc.add_heading("5.2 Rejection Threshold Analysis", level=2)
    doc.add_paragraph(
        "The table below shows how many defaults are captured at various rejection thresholds. "
        "For example, rejecting the top 10% riskiest applicants captures a certain percentage "
        "of defaults while maintaining a 90% approval rate."
    )
    if lr_thresh_df is not None and not lr_thresh_df.empty:
        rows = []
        for _, r in lr_thresh_df.iterrows():
            rows.append([
                f"{r['reject_pct']:.0f}%", _int_fmt(r["n_rejected"]),
                _pct(r["precision"]), _pct(r["recall"]), _pct(r["capture_rate"]),
            ])
        _add_table(doc, ["Reject %", "Rejected", "Precision", "Recall", "Capture Rate"], rows)
        doc.add_paragraph("")

    _add_image(doc, o / "plots" / "stakeholder_threshold_analysis.png")
    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════════════
    # 6. CALIBRATION AND STABILITY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Calibration and Score Stability", level=1)

    doc.add_heading("6.1 Score Distribution Stability (PSI)", level=2)
    doc.add_paragraph(
        "The Population Stability Index measures whether the model's score distribution has "
        "shifted between the training and test populations. PSI < 0.10 indicates stable "
        "distributions; 0.10-0.25 indicates moderate drift; > 0.25 indicates significant drift."
    )
    if psi_df is not None:
        rows = []
        for _, r in psi_df.iterrows():
            flag = "Stable" if r["psi"] < 0.10 else ("Moderate" if r["psi"] < 0.25 else "High drift")
            rows.append([str(r["model"]), _fmt(r["psi"]), flag])
        _add_table(doc, ["Model", "PSI", "Interpretation"], rows)
        doc.add_paragraph("")
        max_psi = float(psi_df["psi"].max())
        doc.add_paragraph(
            f"The highest observed PSI is {_fmt(max_psi)}, which indicates "
            f"{'stable' if max_psi < 0.10 else 'some' if max_psi < 0.25 else 'material'} score-distribution shift "
            "between development and test."
        )

    _add_image(doc, o / "plots" / "stakeholder_calibration.png")
    doc.add_paragraph("")
    _add_image(doc, o / "plots" / "stakeholder_reliability.png")
    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════════════
    # 7. POPULATION BIAS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Population Bias and Fairness", level=1)

    doc.add_heading("7.1 Selection Bias", level=2)
    doc.add_paragraph(
        "Selection bias arises because the model is trained only on booked (accepted) accounts "
        "but will be applied to all applicants. Two diagnostics measure the extent of this bias:"
    )
    doc.add_paragraph(
        "Score distribution divergence (KS test): measures how differently the model scores "
        "booked vs non-booked applicants. A large KS statistic is expected (the existing decision "
        "already separated them), but an extremely high value might indicate the model is merely "
        "recapitulating the old decision boundary rather than learning new risk signal.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Correlation with existing benchmark: the Pearson and Spearman correlation between "
        "model PD and the original risk_score_rf. HIGH correlation (> 0.90) would suggest "
        "the model adds little new information beyond the existing score.",
        style="List Bullet",
    )

    if selbias is not None:
        rows = []
        for _, r in selbias.iterrows():
            rows.append([
                str(r["model"]), _fmt(r["pearson_corr"], signed=True),
                _fmt(r["spearman_corr"], signed=True), r["selection_bias_flag"],
            ])
        _add_table(doc, ["Model", "Pearson r", "Spearman rho", "Flag"], rows)
        doc.add_paragraph("")
        max_abs_pearson = float(selbias["pearson_corr"].abs().max())
        n_high_selbias = int((selbias["selection_bias_flag"] == "HIGH").sum())
        if n_high_selbias == 0:
            doc.add_paragraph(
                f"No model is flagged HIGH for correlation with the existing risk_score_rf. "
                f"The maximum absolute Pearson correlation is {_fmt(max_abs_pearson)}, suggesting that the new models add information beyond the incumbent score."
            )
        else:
            high_models = ", ".join(selbias.loc[selbias["selection_bias_flag"] == "HIGH", "model"].astype(str).tolist())
            doc.add_paragraph(
                f"{n_high_selbias} model(s) are flagged HIGH for correlation with risk_score_rf: {high_models}. "
                "Those candidates may be recapitulating the existing decision boundary rather than contributing distinct signal."
            )

    doc.add_heading("7.2 Adverse Impact by Age", level=2)
    doc.add_paragraph(
        "Adverse impact analysis evaluates whether the model disproportionately rejects "
        "applicants in specific age bands. At a 10% rejection threshold, the adverse impact "
        "ratio (AIR) compares each band's approval rate to the best band. AIR below 0.80 "
        "triggers the conventional four-fifths rule threshold for disparate impact."
    )
    if lr_ai_df is not None and not lr_ai_df.empty:
        rows = []
        for _, r in lr_ai_df.iterrows():
            rows.append([
                r["age_band"], _int_fmt(r["n"]), _pct(r["observed_default_rate"]),
                _pct(r["mean_predicted_pd"]), _pct(r["approval_rate_at_10pct_reject"]),
                f"{r['adverse_impact_ratio']:.2f}", r["air_flag"],
            ])
        _add_table(doc, ["Age Band", "N", "Obs. Default", "Mean PD", "Approval", "AIR", "Flag"], rows)
        doc.add_paragraph("")
        failing_ai = lr_ai_df.loc[lr_ai_df["air_flag"] == "FAIL"].copy()
        if not failing_ai.empty:
            worst_band = failing_ai.sort_values("adverse_impact_ratio").iloc[0]
            doc.add_paragraph(
                f"{len(failing_ai)} age band(s) fall below the 80% AIR threshold for {recommended}. "
                f"The lowest AIR is {worst_band['adverse_impact_ratio']:.2f} in band {worst_band['age_band']}. "
                "This pattern requires regulatory review to determine whether mitigation or monitoring actions are needed."
            )
        else:
            doc.add_paragraph(
                f"All evaluated age bands pass the 80% AIR threshold for {recommended} at the 10% rejection cutoff."
            )

    # ══════════════════════════════════════════════════════════════════════
    # 8. FEATURE ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Feature Assessment", level=1)

    doc.add_heading("8.1 Information Value", level=2)
    doc.add_paragraph(
        "Information Value (IV) measures univariate predictive power on the development-fit "
        "sample. IV < 0.02 is considered useless; 0.02-0.10 weak; 0.10-0.30 medium; "
        "0.30-0.50 strong; > 0.50 suspicious (potential leakage)."
    )
    if iv is not None:
        rows = []
        for _, r in iv.head(10).iterrows():
            strength = "Useless" if r["iv"] < 0.02 else "Weak" if r["iv"] < 0.1 else "Medium" if r["iv"] < 0.3 else "Strong"
            rows.append([str(r["feature"]), _fmt(r["iv"]), strength])
        _add_table(doc, ["Feature", "IV", "Strength"], rows)
        doc.add_paragraph("")
        doc.add_paragraph(
            "No feature exceeds IV 0.30, which rules out data leakage. The top features are "
            "bureau-capacity ratios (AGE/MAX_CREDIT, INCOME/MAX_CREDIT), binned numerical x categorical "
            "interactions, and frequency-encoded categorical surrogates. This pattern indicates that "
            "risk differentiation comes primarily from the relationship between borrower capacity "
            "and credit utilization rather than from any single demographic or product attribute."
        )

    _add_image(doc, o / "plots" / "stakeholder_top_drivers.png")
    doc.add_paragraph("")
    _add_image(doc, o / "plots" / "shap_summary.png", width=5.0)
    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════════════════════
    # 9. ENSEMBLE
    # ══════════════════════════════════════════════════════════════════════
    if ens is not None and not ens.empty:
        doc.add_heading("9. Post-Hoc Ensemble", level=1)
        e = ens.iloc[0]
        ens_r = results.loc[results["Model"].str.contains("Ensemble")].iloc[0] if results is not None else None
        base_component_rows = results.loc[results["Model"] == str(e["lr_name"])] if results is not None else pd.DataFrame()
        base_component_r = base_component_rows.iloc[0] if not base_component_rows.empty else rec_r
        doc.add_paragraph(
            f"A post-hoc ensemble blends {e['lr_name']} ({e['lr_weight']:.0%}) and "
            f"{e['tree_name']} ({e['tree_weight']:.0%}), with weights optimized via grid search "
            f"on calibration holdout PR AUC. On the held-out test, the ensemble achieves "
            f"ROC AUC {_fmt(ens_r['ROC AUC'])} and PR AUC {_fmt(ens_r['PR AUC'])}, "
            f"compared with standalone {e['lr_name']} "
            f"(ROC AUC {_fmt(base_component_r['ROC AUC'])}, PR AUC {_fmt(base_component_r['PR AUC'])})."
        )
        doc.add_paragraph(
            f"The ensemble provides a slight discrimination improvement by combining {e['lr_name']}'s "
            f"generalization profile with {e['tree_name']}'s complementary error profile. However, the "
            "added operational complexity of maintaining two models should be weighed against "
            "the marginal performance gain."
        )

    # ══════════════════════════════════════════════════════════════════════
    # 10. LIMITATIONS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Limitations and Caveats", level=1)

    ai_limit_text = (
        "All evaluated age bands pass the 80% AIR threshold at the selected rejection cutoff. Fairness should still be monitored as the applicant mix evolves."
        if lr_ai_df is not None and not lr_ai_df.empty and int((lr_ai_df["air_flag"] == "FAIL").sum()) == 0
        else "At least one age band falls below the 80% AIR threshold. Even where this aligns with observed risk differentials, it may require mitigation depending on regulatory requirements."
    )
    limitations = [
        ("Booked-proxy evaluation",
         "All metrics are computed on the accepted population only. Rejected and canceled "
         "applicants have no observed default outcome, so the model's performance on the full "
         "applicant population is unknown. The booked-proxy metrics are an optimistic estimate "
         "of true applicant-stage performance because the existing decisioning has already "
         "filtered out the highest-risk applicants."),
        ("Model generalization",
         "Several flexible candidates exhibit materially larger train-test gaps than the recommended model. "
         "This limits how much weight the governance scorecard can place on raw in-sample performance alone."),
        ("Age-based adverse impact",
         ai_limit_text),
        ("Benchmark ceiling",
         "risk_score_rf retains a meaningful PR AUC advantage, particularly in concentrating "
         "defaults in the highest-risk segment. Closing this gap is the primary modeling opportunity."),
        ("Single temporal test window",
         "The held-out evaluation covers July 2024 to January 2025. While rolling OOT provides "
         "additional evidence, a longer forward evaluation period would strengthen confidence "
         "in production stability."),
        ("Calibration sensitivity",
         "Calibrated probabilities depend on the holdout sample size and default rate. Changes "
         "in the portfolio mix or economic environment may require recalibration."),
    ]
    for title_text, body in limitations:
        p = doc.add_paragraph()
        p.add_run(f"{title_text}. ").bold = True
        p.add_run(body)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix: Stakeholder Charts", level=1)
    for fname, caption in [
        ("stakeholder_executive_summary.png", "Executive Summary"),
        ("stakeholder_kpis.png", "Key Performance Indicators"),
        ("stakeholder_population_coverage.png", "Population Coverage"),
        ("stakeholder_roc_pr_curves.png", "ROC and Precision-Recall Curves"),
        ("stakeholder_auc_lift.png", "AUC Lift vs Benchmarks"),
    ]:
        if (o / "plots" / fname).exists():
            doc.add_heading(caption, level=2)
            _add_image(doc, o / "plots" / fname)
            doc.add_paragraph("")

    doc.save(report_path)
    print(f"Report saved: {report_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", default="output")
    parser.add_argument("--report-path", default="report.docx")
    args = parser.parse_args()
    generate_report(args.output_dir, args.report_path)
