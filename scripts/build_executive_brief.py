"""Build a credit-committee executive brief (.docx) — fits on two pages.

Audience: credit committee, business sponsor. The methodology report
(scripts/build_methodology_report.py) is the long-form artifact; this brief
is for a 5-minute decision.

Usage:
    uv run python scripts/build_executive_brief.py
    uv run python scripts/build_executive_brief.py --output executive_brief.docx
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _read_csv(path: Path) -> pd.DataFrame | None:
    return pd.read_csv(path) if path.exists() else None


def _fmt(val, digits: int = 4, signed: bool = False) -> str:
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "n/a"
    sign = "+" if signed else ""
    return f"{float(val):{sign}.{digits}f}"


def _set_cell_bg(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_borders(table, color: str = "BBBBBB") -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        borders.append(b)
    tblPr.append(borders)


def _kpi_tile(table, row_idx: int, col_idx: int, label: str, value: str, sub: str | None = None) -> None:
    cell = table.rows[row_idx].cells[col_idx]
    _set_cell_bg(cell, "F1F5F9")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p_label = cell.paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_label.add_run(label)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x65, 0x75)
    p_value = cell.add_paragraph()
    p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = p_value.add_run(value)
    rv.bold = True
    rv.font.size = Pt(20)
    rv.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    if sub:
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = p_sub.add_run(sub)
        rs.italic = True
        rs.font.size = Pt(8)
        rs.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_para(doc, text: str, *, size: int = 11, bold: bool = False, italic: bool = False, color: tuple[int, int, int] | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = RGBColor(*color)


def _verdict_banner(doc: Document, *, deploy_ok: bool, recommended: str, score: float, floor: float = 50.0) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    color = "16A34A" if deploy_ok else "DC2626"
    _set_cell_bg(cell, color)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline = (
        f"DEPLOY: {recommended} (calibrated)"
        if deploy_ok else "DO NOT DEPLOY — no candidate cleared the quality floor"
    )
    r1 = p1.add_run(headline)
    r1.bold = True
    r1.font.size = Pt(18)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = (
        f"Weighted selection score {score:.1f}/100 (floor {floor:.0f})"
        if deploy_ok
        else f"Best candidate {recommended} scored {score:.1f}/100 (floor {floor:.0f})"
    )
    r2 = p2.add_run(sub)
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def build_brief(output_dir: Path, brief_path: Path) -> None:
    results = _read_csv(output_dir / "results.csv")
    if results is not None:
        results = results.set_index(results.columns[0])
    sel = _read_csv(output_dir / "model_selection.csv")
    bench = _read_csv(output_dir / "benchmark_comparisons.csv")
    cis = _read_csv(output_dir / "confidence_intervals.csv")
    if cis is not None and "Model" in cis.columns:
        cis = cis.set_index("Model")
    cal = _read_csv(output_dir / "calibration_diagnostics.csv")
    psi = _read_csv(output_dir / "psi.csv")
    air = _read_csv(output_dir / "adverse_impact_age.csv")
    rolling = _read_csv(output_dir / "rolling_oot_summary.csv")
    plots_dir = output_dir / "plots"

    # Pull recommendation
    rec = "Logistic Regression"
    score = float("nan")
    deploy_ok = False
    if sel is not None and not sel.empty:
        winners = sel.loc[sel.get("recommended", False).fillna(False).astype(bool)]
        if not winners.empty:
            rec = str(winners.iloc[0]["model"])
            score = float(winners.iloc[0]["weighted_score"])
            deploy_ok = bool(winners.iloc[0].get("meets_quality_floor", True))
        else:
            sel_sorted = sel.sort_values("weighted_score", ascending=False)
            rec = str(sel_sorted.iloc[0]["model"])
            score = float(sel_sorted.iloc[0]["weighted_score"])
            deploy_ok = False

    rec_cal = f"{rec} (calibrated)"
    rec_res = results.loc[rec] if (results is not None and rec in results.index) else None
    rec_cal_res = results.loc[rec_cal] if (results is not None and rec_cal in results.index) else None

    # ── Document ─────────────────────────────────────────────────────────────
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tighten margins so we fit on one page when possible
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Title strip
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t = title_p.add_run("basel_bad PD model — Executive Brief")
    t.bold = True
    t.font.size = Pt(20)
    t.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    sub_p = doc.add_paragraph()
    s = sub_p.add_run(f"Credit committee summary · {date.today().isoformat()}")
    s.italic = True
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Verdict banner
    _verdict_banner(doc, deploy_ok=deploy_ok, recommended=rec, score=score)

    # ── KPI grid (2 rows × 4 cols) ───────────────────────────────────────────
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = 1
    for row in kpi_table.rows:
        for cell in row.cells:
            cell.width = Inches(1.65)

    # Row 1: discrimination
    auc = float(rec_res.get("ROC AUC")) if rec_res is not None else float("nan")
    pr_auc = float(rec_res.get("PR AUC")) if rec_res is not None else float("nan")
    ks = float(rec_res.get("KS")) if rec_res is not None else float("nan")
    n = int(rec_res.get("N", 0)) if rec_res is not None else 0

    auc_ci = ""
    pr_ci = ""
    if cis is not None and rec in cis.index:
        auc_ci = f"95% CI [{float(cis.loc[rec, 'AUC_lo']):.3f}, {float(cis.loc[rec, 'AUC_hi']):.3f}]"
        pr_ci = f"95% CI [{float(cis.loc[rec, 'PR_AUC_lo']):.3f}, {float(cis.loc[rec, 'PR_AUC_hi']):.3f}]"

    _kpi_tile(kpi_table, 0, 0, "ROC AUC", _fmt(auc, 4), auc_ci)
    _kpi_tile(kpi_table, 0, 1, "Gini", _fmt(2 * auc - 1, 4) if not np.isnan(auc) else "n/a")
    _kpi_tile(kpi_table, 0, 2, "KS statistic", _fmt(ks, 4))
    _kpi_tile(kpi_table, 0, 3, "PR AUC", _fmt(pr_auc, 4), pr_ci)

    # Row 2: calibration + ECE + capture + n
    brier = float(rec_cal_res.get("Brier")) if rec_cal_res is not None and not pd.isna(rec_cal_res.get("Brier")) else float("nan")
    ece = float("nan")
    if cal is not None and not cal.empty:
        cal_row = cal.loc[cal["model"] == rec_cal]
        if not cal_row.empty:
            ece = float(cal_row.iloc[0]["ece"])
    # Capture at top 10% comes from the lift table; if absent, leave n/a
    lift_path = output_dir / "lift_table.csv"
    capture_10 = float("nan")
    if lift_path.exists():
        lift = pd.read_csv(lift_path)
        rec_lift = lift.loc[lift["model"] == rec]
        if not rec_lift.empty and "decile" in rec_lift.columns and "capture_rate" in rec_lift.columns:
            top = rec_lift.loc[rec_lift["decile"] == 1]
            if not top.empty:
                capture_10 = float(top.iloc[0]["capture_rate"])

    _kpi_tile(kpi_table, 1, 0, "Calibrated Brier", _fmt(brier, 4), "lower = better")
    _kpi_tile(kpi_table, 1, 1, "ECE", _fmt(ece, 4) if not np.isnan(ece) else "n/a", "calibration error")
    _kpi_tile(kpi_table, 1, 2, "Top-10% capture", f"{capture_10:.1%}" if not np.isnan(capture_10) else "n/a")
    _kpi_tile(kpi_table, 1, 3, "Test sample", f"{n:,}" if n else "n/a", "booked-proxy rows")
    _set_borders(kpi_table)

    # ── Recommendation rationale ────────────────────────────────────────────
    doc.add_paragraph()
    rationale_p = doc.add_paragraph()
    rh = rationale_p.add_run("Recommendation rationale")
    rh.bold = True
    rh.font.size = Pt(12)
    rh.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    rationale_lines: list[str] = []
    if sel is not None and not sel.empty and "meets_quality_floor" in sel.columns:
        n_clear = int(sel["meets_quality_floor"].sum())
        n_total = len(sel)
        rationale_lines.append(f"{n_clear} of {n_total} candidates clear the absolute weighted-score floor of 50.")
    if bench is not None and not bench.empty:
        cand_rows = bench.loc[bench["candidate_model"] == rec]
        for ref in ("score_RF (benchmark)", "risk_score_rf (benchmark)"):
            r = cand_rows.loc[cand_rows["reference_model"] == ref]
            if not r.empty:
                row = r.iloc[0]
                imp = float(row["auc_improvement"])
                p_adj = float(row["auc_p_adjusted"])
                lo = float(row["auc_improvement_lo"])
                outcome = "win" if (imp > 0 and lo > 0) else ("statistical tie" if imp >= 0 else "loss")
                rationale_lines.append(
                    f"vs. {ref}: ΔAUC = {imp:+.4f} (lower CI {lo:+.4f}, adj. p = {p_adj:.3f}) — {outcome}."
                )
    if rolling is not None and not rolling.empty:
        rec_rows = rolling.loc[rolling["Model"] == rec_cal]
        if rec_rows.empty:
            rec_rows = rolling.loc[rolling["Model"] == rec]
        if not rec_rows.empty:
            row = rec_rows.iloc[0]
            rationale_lines.append(
                f"Rolling-OOT stability: PR AUC = {float(row['mean_PR_AUC']):.4f} "
                f"(±{float(row['std_PR_AUC']):.4f}) across {int(row['n_folds'])} folds."
            )
    for line in rationale_lines[:5]:
        p = doc.add_paragraph(line, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # ── Watch-list ──────────────────────────────────────────────────────────
    watch_p = doc.add_paragraph()
    wh = watch_p.add_run("Watch-list (issues for monitoring or follow-up)")
    wh.bold = True
    wh.font.size = Pt(12)
    wh.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

    watch_items: list[str] = []
    # PSI high drift
    if psi is not None and not psi.empty:
        high = psi.loc[psi["psi"] >= 0.25]
        if not high.empty:
            names = ", ".join(high["model"].astype(str).head(3))
            watch_items.append(f"PSI high drift on score distributions: {names}.")
    # Adverse impact failures
    if air is not None and not air.empty and "ratio_status" in air.columns:
        fails = air.loc[air["ratio_status"] == "FAIL"]
        if not fails.empty and "model" in fails.columns and "age_band" in fails.columns:
            bands_by_model = (
                fails.groupby("model")["age_band"].apply(lambda b: ", ".join(map(str, b.unique())))
                .to_dict()
            )
            sample = "; ".join(f"{m}: {b}" for m, b in list(bands_by_model.items())[:3])
            watch_items.append(f"Adverse-impact ratios below 0.80 — {sample}.")
    # Benchmark gap
    if bench is not None and not bench.empty:
        rec_vs_rf = bench.loc[(bench["candidate_model"] == rec) & (bench["reference_model"] == "risk_score_rf (benchmark)")]
        if not rec_vs_rf.empty:
            pr_imp = float(rec_vs_rf.iloc[0]["pr_auc_improvement"])
            if pr_imp < 0:
                watch_items.append(
                    f"Trained model loses to risk_score_rf on PR AUC by {abs(pr_imp):.4f} — "
                    "feature ceiling vs. existing score (see methodology report §13)."
                )
    if not watch_items:
        watch_items.append("No watch-list items flagged in this run. Continue standard monitoring cadence.")

    for line in watch_items:
        p = doc.add_paragraph(line, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # ── Chart embed (small, single) ─────────────────────────────────────────
    chart = plots_dir / "stakeholder_executive_summary.png"
    if chart.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(chart), width=Inches(6.5))

    # ── Footer pointer ──────────────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph()
    f = foot.add_run(
        "Full methodology + results are documented in methodology_report.docx. "
        "Production deployment requires the SHA-256 model checksums in output/models/checksums.json."
    )
    f.italic = True
    f.font.size = Pt(9)
    f.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.save(str(brief_path))
    print(f"Saved executive brief: {brief_path}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", default="output", help="Pipeline output directory")
    parser.add_argument("--brief", default="executive_brief.docx", help="Path to write the .docx brief")
    args = parser.parse_args(argv)
    build_brief(Path(args.output_dir), Path(args.brief))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
